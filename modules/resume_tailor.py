"""Tailors a master profile to a specific job description using Claude AI.

Outputs:
  - ATS-optimised .docx  (no tables, no columns, no graphics)
  - ATS-optimised .pdf   (reportlab — no Word dependency)
  - Match score 0-100 with keyword breakdown
"""
import json
import re
import textwrap
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

OUTPUT_DIR = Path(__file__).parent.parent / "data" / "resumes"
MODEL = "claude-sonnet-4-6"

SYSTEM_PROMPT = textwrap.dedent("""\
    You are a world-class ATS resume writer and recruiter with 15 years of experience.
    Your task is to tailor a candidate's master profile to a specific job description.

    STRICT ATS RULES — never violate these:
    - No tables, columns, text boxes, headers/footers, or graphics of any kind
    - Standard section headings only: PROFESSIONAL SUMMARY, EXPERIENCE, EDUCATION,
      SKILLS, PROJECTS, CERTIFICATIONS
    - Bullet points start with strong action verbs
    - Every achievement bullet must include at least one quantified metric where possible
    - Maximum 7 bullets per job; minimum 3
    - Skills section shows ONLY skills relevant to this specific job description
    - Summary must mirror the exact language and keywords in the job description
    - Keep all facts 100% truthful — do not invent experience or credentials
    - Rewrite bullets to naturally embed exact keywords from the job description

    MATCH SCORE RULES:
    - keyword_match (0-100): % of important JD keywords present in the tailored resume
    - skills_alignment (0-100): % of required/preferred skills the candidate has
    - experience_relevance (0-100): how closely experience matches JD seniority and domain
    - overall (0-100): weighted average (40% keyword, 35% skills, 25% experience)

    Respond with ONLY a JSON object inside a ```json code block. No other text.
    The JSON must follow this exact schema:
    {
      "summary": "<2-3 sentence professional summary>",
      "experience": [
        {
          "company": "...",
          "title": "...",
          "start_date": "...",
          "end_date": "...",
          "location": "...",
          "bullets": ["...", "..."]
        }
      ],
      "education": [
        {
          "degree": "...",
          "institution": "...",
          "graduation_year": "...",
          "gpa": "...",
          "relevant_courses": ["..."]
        }
      ],
      "skills": {
        "technical": ["..."],
        "soft": ["..."],
        "tools": ["..."]
      },
      "projects": [
        {
          "name": "...",
          "description": "...",
          "tech_stack": ["..."],
          "live_url": "...",
          "github_url": "..."
        }
      ],
      "certifications": [
        {
          "name": "...",
          "issuer": "...",
          "date": "..."
        }
      ],
      "match_score": {
        "overall": 0,
        "keyword_match": 0,
        "skills_alignment": 0,
        "experience_relevance": 0
      },
      "keywords_found": ["..."],
      "keywords_missing": ["..."]
    }
""")


class ResumeTailor:
    def __init__(self, api_key: str):
        self.client = anthropic.Anthropic(api_key=api_key)

    # ---------------------------------------------------------------- Public --

    def tailor(self, profile: dict, job_description: str, company: str, job_title: str) -> dict:
        """Tailor the profile to the job description.

        Returns:
            {
                "tailored":     dict,       # the Claude-generated resume sections
                "match_score":  dict,       # overall + breakdown scores
                "keywords_found":  list,
                "keywords_missing": list,
                "docx_path":    str,
                "pdf_path":     str,
            }
        """
        prompt = self._build_prompt(profile, job_description)
        raw = self._call_claude(prompt)
        tailored = self._parse_response(raw)

        slug = _make_slug(company, job_title)
        out_dir = OUTPUT_DIR / slug
        out_dir.mkdir(parents=True, exist_ok=True)

        docx_path = out_dir / "resume.docx"
        pdf_path = out_dir / "resume.pdf"

        self._save_docx(tailored, profile, str(docx_path))
        self._save_pdf(tailored, profile, str(pdf_path))

        return {
            "tailored": tailored,
            "match_score": tailored.pop("match_score", {}),
            "keywords_found": tailored.pop("keywords_found", []),
            "keywords_missing": tailored.pop("keywords_missing", []),
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path),
        }

    # --------------------------------------------------------------- Claude --

    def _build_prompt(self, profile: dict, job_description: str) -> str:
        return (
            "## CANDIDATE MASTER PROFILE\n"
            f"{json.dumps(profile, indent=2)}\n\n"
            "## JOB DESCRIPTION\n"
            f"{job_description}\n\n"
            "Tailor the resume to this job description following all rules in the system prompt."
        )

    def _call_claude(self, prompt: str) -> str:
        message = self.client.messages.create(
            model=MODEL,
            max_tokens=4096,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text

    def _parse_response(self, raw: str) -> dict:
        # Extract JSON from ```json ... ``` block
        match = re.search(r"```json\s*([\s\S]*?)\s*```", raw)
        if match:
            json_str = match.group(1)
        else:
            # Fall back: find outermost { ... }
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start == -1 or end == 0:
                raise ValueError("Claude response contained no parseable JSON.")
            json_str = raw[start:end]

        try:
            return json.loads(json_str)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Failed to parse Claude JSON response: {exc}") from exc

    # ----------------------------------------------------------------- DOCX --

    def _save_docx(self, content: dict, profile: dict, path: str) -> None:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Default font for Normal style
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        personal = profile["personal"]

        # ── Name ────────────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(personal.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(16)

        # ── Contact line ────────────────────────────────────────────────────
        contact_parts = [
            personal.get("email", ""),
            personal.get("phone", ""),
            personal.get("location", ""),
            personal.get("linkedin_url", ""),
            personal.get("github_url", ""),
        ]
        contact_line = "  |  ".join(p for p in contact_parts if p)
        cp = doc.add_paragraph(contact_line)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(9.5)

        # ── Summary ─────────────────────────────────────────────────────────
        if content.get("summary"):
            _docx_section_header(doc, "PROFESSIONAL SUMMARY")
            sp = doc.add_paragraph(content["summary"])
            sp.paragraph_format.space_after = Pt(4)

        # ── Experience ──────────────────────────────────────────────────────
        if content.get("experience"):
            _docx_section_header(doc, "EXPERIENCE")
            for job in content["experience"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(f"{job.get('title', '')}  |  {job.get('company', '')}")
                r.bold = True
                r.font.size = Pt(11)

                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(2)
                mr = mp.add_run(
                    f"{job.get('location', '')}   •   "
                    f"{job.get('start_date', '')} – {job.get('end_date', '')}"
                )
                mr.italic = True
                mr.font.size = Pt(9.5)

                for bullet in job.get("bullets", []):
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(1)
                    bp.add_run(bullet)

        # ── Education ───────────────────────────────────────────────────────
        if content.get("education"):
            _docx_section_header(doc, "EDUCATION")
            for edu in content["education"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(f"{edu.get('degree', '')}  |  {edu.get('institution', '')}")
                r.bold = True

                details = []
                if edu.get("graduation_year"):
                    details.append(f"Graduated {edu['graduation_year']}")
                if edu.get("gpa"):
                    details.append(f"GPA: {edu['gpa']}")
                if details:
                    dp = doc.add_paragraph("  |  ".join(details))
                    dp.paragraph_format.space_after = Pt(1)
                    for r in dp.runs:
                        r.italic = True
                        r.font.size = Pt(9.5)

                if edu.get("relevant_courses"):
                    cp = doc.add_paragraph()
                    cp.paragraph_format.space_after = Pt(2)
                    rb = cp.add_run("Relevant Courses: ")
                    rb.bold = True
                    cp.add_run(", ".join(edu["relevant_courses"]))

        # ── Skills ──────────────────────────────────────────────────────────
        skills = content.get("skills", {})
        if any(skills.get(k) for k in ("technical", "tools", "soft")):
            _docx_section_header(doc, "SKILLS")
            for label, key in [("Technical", "technical"), ("Tools", "tools"), ("Soft Skills", "soft")]:
                items = skills.get(key, [])
                if items:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    rb = p.add_run(f"{label}: ")
                    rb.bold = True
                    p.add_run(", ".join(items))

        # ── Projects ────────────────────────────────────────────────────────
        if content.get("projects"):
            _docx_section_header(doc, "PROJECTS")
            for proj in content["projects"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(proj.get("name", ""))
                r.bold = True
                if proj.get("tech_stack"):
                    p.add_run(f"  |  {', '.join(proj['tech_stack'])}")

                dp = doc.add_paragraph(proj.get("description", ""))
                dp.paragraph_format.space_after = Pt(1)

                links = []
                if proj.get("live_url"):
                    links.append(f"Live: {proj['live_url']}")
                if proj.get("github_url"):
                    links.append(f"GitHub: {proj['github_url']}")
                if links:
                    lp = doc.add_paragraph("  |  ".join(links))
                    lp.paragraph_format.space_after = Pt(2)
                    for r in lp.runs:
                        r.font.size = Pt(9.5)

        # ── Certifications ───────────────────────────────────────────────────
        if content.get("certifications"):
            _docx_section_header(doc, "CERTIFICATIONS")
            for cert in content["certifications"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(cert.get("name", ""))
                r.bold = True
                p.add_run(f"  |  {cert.get('issuer', '')}  |  {cert.get('date', '')}")

        doc.save(path)

    # ------------------------------------------------------------------ PDF --

    def _save_pdf(self, content: dict, profile: dict, path: str) -> None:
        doc = SimpleDocTemplate(
            path,
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # ── Styles ──────────────────────────────────────────────────────────
        S = _pdf_styles()
        story = []
        personal = profile["personal"]

        # Name
        story.append(Paragraph(personal.get("full_name", ""), S["name"]))

        # Contact
        contact_parts = [
            personal.get("email", ""),
            personal.get("phone", ""),
            personal.get("location", ""),
            personal.get("linkedin_url", ""),
            personal.get("github_url", ""),
        ]
        contact_line = "  |  ".join(p for p in contact_parts if p)
        story.append(Paragraph(contact_line, S["contact"]))
        story.append(Spacer(1, 4))

        # Summary
        if content.get("summary"):
            story += _pdf_section("PROFESSIONAL SUMMARY", S)
            story.append(Paragraph(content["summary"], S["body"]))

        # Experience
        if content.get("experience"):
            story += _pdf_section("EXPERIENCE", S)
            for job in content["experience"]:
                story.append(
                    Paragraph(
                        f"<b>{job.get('title', '')}  |  {job.get('company', '')}</b>",
                        S["job_title"],
                    )
                )
                story.append(
                    Paragraph(
                        f"<i>{job.get('location', '')}   •   "
                        f"{job.get('start_date', '')} – {job.get('end_date', '')}</i>",
                        S["job_meta"],
                    )
                )
                for bullet in job.get("bullets", []):
                    story.append(Paragraph(f"• {bullet}", S["bullet"]))
                story.append(Spacer(1, 4))

        # Education
        if content.get("education"):
            story += _pdf_section("EDUCATION", S)
            for edu in content["education"]:
                story.append(
                    Paragraph(
                        f"<b>{edu.get('degree', '')}  |  {edu.get('institution', '')}</b>",
                        S["job_title"],
                    )
                )
                details = []
                if edu.get("graduation_year"):
                    details.append(f"Graduated {edu['graduation_year']}")
                if edu.get("gpa"):
                    details.append(f"GPA: {edu['gpa']}")
                if details:
                    story.append(Paragraph(f"<i>{'  |  '.join(details)}</i>", S["job_meta"]))
                if edu.get("relevant_courses"):
                    story.append(
                        Paragraph(
                            f"<b>Relevant Courses:</b> {', '.join(edu['relevant_courses'])}",
                            S["body"],
                        )
                    )
                story.append(Spacer(1, 4))

        # Skills
        skills = content.get("skills", {})
        if any(skills.get(k) for k in ("technical", "tools", "soft")):
            story += _pdf_section("SKILLS", S)
            for label, key in [("Technical", "technical"), ("Tools", "tools"), ("Soft Skills", "soft")]:
                items = skills.get(key, [])
                if items:
                    story.append(
                        Paragraph(f"<b>{label}:</b> {', '.join(items)}", S["body"])
                    )

        # Projects
        if content.get("projects"):
            story += _pdf_section("PROJECTS", S)
            for proj in content["projects"]:
                tech = f"  |  {', '.join(proj['tech_stack'])}" if proj.get("tech_stack") else ""
                story.append(
                    Paragraph(f"<b>{proj.get('name', '')}</b>{tech}", S["job_title"])
                )
                story.append(Paragraph(proj.get("description", ""), S["body"]))
                links = []
                if proj.get("live_url"):
                    links.append(f"Live: {proj['live_url']}")
                if proj.get("github_url"):
                    links.append(f"GitHub: {proj['github_url']}")
                if links:
                    story.append(Paragraph("  |  ".join(links), S["job_meta"]))
                story.append(Spacer(1, 4))

        # Certifications
        if content.get("certifications"):
            story += _pdf_section("CERTIFICATIONS", S)
            for cert in content["certifications"]:
                story.append(
                    Paragraph(
                        f"<b>{cert.get('name', '')}</b>  |  "
                        f"{cert.get('issuer', '')}  |  {cert.get('date', '')}",
                        S["body"],
                    )
                )

        doc.build(story)


# ─────────────────────────────────────── DOCX helpers ────────────────────────

def _docx_section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

    # Bottom border under the heading paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────── PDF helpers ─────────────────────────

def _pdf_styles() -> dict:
    return {
        "name": ParagraphStyle(
            "name", fontSize=16, fontName="Helvetica-Bold",
            alignment=TA_CENTER, spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "contact", fontSize=9, fontName="Helvetica",
            alignment=TA_CENTER, spaceAfter=6,
        ),
        "section_head": ParagraphStyle(
            "section_head", fontSize=11, fontName="Helvetica-Bold",
            spaceBefore=10, spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "body", fontSize=10, fontName="Helvetica",
            spaceAfter=3, leading=14,
        ),
        "bullet": ParagraphStyle(
            "bullet", fontSize=10, fontName="Helvetica",
            leftIndent=14, firstLineIndent=-10,
            spaceAfter=2, leading=13,
        ),
        "job_title": ParagraphStyle(
            "job_title", fontSize=10.5, fontName="Helvetica-Bold",
            spaceBefore=6, spaceAfter=1,
        ),
        "job_meta": ParagraphStyle(
            "job_meta", fontSize=9.5, fontName="Helvetica-Oblique",
            spaceAfter=3,
        ),
    }


def _pdf_section(title: str, S: dict) -> list:
    return [
        HRFlowable(width="100%", thickness=0.75, color=colors.black, spaceAfter=2),
        Paragraph(title, S["section_head"]),
    ]


# ──────────────────────────────────────────── Util ───────────────────────────

def _make_slug(company: str, job_title: str) -> str:
    raw = f"{company}_{job_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return re.sub(r"[^a-zA-Z0-9_-]", "_", raw)[:80]
